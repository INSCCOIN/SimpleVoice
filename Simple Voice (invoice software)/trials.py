## ICSS LLC 2024


import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QDialog, QHBoxLayout, QListWidget, QComboBox, QFileDialog, QFontDialog, QTableWidget, QTableWidgetItem, QFormLayout, QInputDialog, QVBoxLayout, QWidget, QLabel, QLineEdit, QPushButton, QGridLayout, QMessageBox
from PyQt5.QtCore import Qt
import pandas as pd
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import rsa
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.asymmetric import padding




# JSON database for INVOICES

def load_invoices():
    try:
        with open('INVOICES.json', 'r') as file:
            return json.load(file)
    except (FileNotFoundError, json.JSONDecodeError):
        return []  # Return an empty list if the file doesn't exist or is empty

def save_invoices(invoices):
    with open('INVOICES.json', 'w') as file:
        json.dump(invoices, file, indent=4)






## SECURITY KEY GEN / CHECK

def generate_keys():
    private_key = rsa.generate_private_key(
        backend=default_backend(),
        public_exponent=65537,
        key_size=2048
    )
    public_key = private_key.public_key()
    # Save the private and public keys to files


def sign_invoice(invoice_content, private_key):
    signature = private_key.sign(
        invoice_content.encode(),
        padding.PSS(
            mgf=padding.MGF1(hashes.SHA256()),
            salt_length=padding.PSS.MAX_LENGTH
        ),
        hashes.SHA256()
    )
    return signature

def verify_invoice(invoice_content, signature, public_key):
    try:
        public_key.verify(
            signature,
            invoice_content.encode(),
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),
                salt_length=padding.PSS.MAX_LENGTH
            ),
            hashes.SHA256()
        )
        return True
    except:
        return False

# SECURITY KEY GEN / CHECK






class InvoiceApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.invoice_items = []
        self.invoices = {}  # Stores all invoices
        self.current_invoice_id = 0  # Unique identifier for each invoice
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Invoice Generator')
        self.setGeometry(100, 100, 800, 600)

        # Central widget
        central_widget = QWidget(self)
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # Create UI elements and layout
        self.create_menu_bar()
        self.create_form_layout(main_layout)
        self.create_invoice_management_layout(main_layout)
        self.create_bottom_buttons(main_layout)
        self.create_invoice_table(main_layout)

        # Refresh invoice list
        self.refresh_invoice_list()

    def create_menu_bar(self):
        self.menu_bar = self.menuBar()
        self.settings_menu = self.menu_bar.addMenu('Settings')
        self.settings_action = self.settings_menu.addAction('Customize Invoice')
        self.settings_action.triggered.connect(self.open_settings)

    def create_form_layout(self, layout):
        form_layout = QFormLayout()
        # Add all your QLineEdit and QPushButton widgets to form_layout
        # ...
        layout.addLayout(form_layout)

    def create_invoice_management_layout(self, layout):
        invoice_management_layout = QHBoxLayout()

        # Invoice List Widget
        self.invoice_list_widget = QListWidget()
        invoice_management_layout.addWidget(self.invoice_list_widget)

        # Add Invoice Button
        self.add_invoice_button = QPushButton('Add Invoice')
        self.add_invoice_button.clicked.connect(self.add_invoice)
        invoice_management_layout.addWidget(self.add_invoice_button)

        # Edit Invoice Button
        self.edit_invoice_button = QPushButton('Edit Invoice')
        self.edit_invoice_button.clicked.connect(self.edit_invoice)
        invoice_management_layout.addWidget(self.edit_invoice_button)

        # Export Invoice Button
        self.export_invoice_button = QPushButton('Export Invoice')
        self.export_invoice_button.clicked.connect(self.export_invoice)
        invoice_management_layout.addWidget(self.export_invoice_button)

        # Add the entire invoice management layout to the main layout
        layout.addLayout(invoice_management_layout)
        
    def add_invoice(self):
        # Collect invoice data
        new_invoice_data = self.collect_invoice_data()

        # Generate a unique invoice ID
        new_invoice_id = self.generate_new_invoice_id()
        new_invoice_data['id'] = new_invoice_id

        # Add the new invoice to the invoices dictionary
        self.invoices[new_invoice_id] = new_invoice_data

        # Save the updated invoices list
        save_invoices(self.invoices)

        # Refresh the invoice list in the UI
        self.refresh_invoice_list()
        
    def edit_invoice(self):
        selected_invoice_id = self.get_selected_invoice_id()
        if selected_invoice_id is not None:
            invoice_data = self.invoices.get(selected_invoice_id, None)
            if invoice_data:
                # Open a dialog to edit the invoice
                edit_dialog = EditInvoiceDialog(invoice_data, self)
                if edit_dialog.exec_():
                    # Update the invoice data
                    updated_invoice_data = edit_dialog.get_updated_data()
                    self.invoices[selected_invoice_id] = updated_invoice_data

                    # Save and refresh
                    save_invoices(self.invoices)
                    self.refresh_invoice_list()
            
    def export_invoice(self):
        selected_invoice_id = self.get_selected_invoice_id()
        if selected_invoice_id is not None:
            invoice_data = self.invoices.get(selected_invoice_id, None)
            if invoice_data:
                # Export the invoice data
                # Example: Exporting to PDF (you need to implement the export_to_pdf function)
                export_to_excel(invoice_data)

                # Inform the user that the invoice was exported
                QMessageBox.information(self, "Export", "Invoice exported successfully.")  

    def create_bottom_buttons(self, layout):
        bottom_button_layout = QHBoxLayout()

        self.generate_button = QPushButton('Generate Invoice', self)
        self.generate_button.clicked.connect(self.generate_invoice)

        self.export_button = QPushButton('Export Invoice', self)
        self.export_button.clicked.connect(self.export_invoice)

        self.create_invoice_button = QPushButton("Create Invoice", self)
        self.create_invoice_button.clicked.connect(self.on_create_invoice_clicked)

        bottom_button_layout.addWidget(self.generate_button)
        bottom_button_layout.addWidget(self.export_button)
        bottom_button_layout.addWidget(self.create_invoice_button)

        layout.addLayout(bottom_button_layout)

    def create_invoice_table(self, layout):
        self.invoice_table = QTableWidget()
        self.setup_invoice_table()
        layout.addWidget(self.invoice_table)

    def setup_invoice_table(self):
        self.invoice_table.setColumnCount(3)  # Adjust the number of columns as needed
        self.invoice_table.setHorizontalHeaderLabels(["Invoice ID", "Client Name", "Amount"])

    def setup_invoice_table(self):
        # Set up columns, headers, etc. for invoice_table
        self.invoice_table.setColumnCount(3)  # Example: 3 columns
        self.invoice_table.setHorizontalHeaderLabels(["Invoice ID", "Client Name", "Amount"])

    def create_status_dropdown(self, invoice_id):
        dropdown = QComboBox(self)
        dropdown.addItem("Paid")
        dropdown.addItem("Unpaid")
        dropdown.addItem("Overdue")
        dropdown.currentIndexChanged.connect(lambda: self.on_status_changed(invoice_id, dropdown.currentText()))
        return dropdown

    def on_status_changed(self, invoice_id, new_status):
        self.update_invoice_status(invoice_id, new_status)

    def on_create_invoice_clicked(self):
        # Assuming you collect invoice data into a variable named 'new_invoice_data'
        new_invoice_data = self.collect_invoice_data()
        self.add_invoice(new_invoice_data)

    def add_invoice(self, new_invoice):
        invoices = load_invoices()
        invoices.append(new_invoice)
        save_invoices(invoices)
        self.refresh_invoice_list()  # Refresh the invoice list in the UI

    def update_invoice_status(self, invoice_id, new_status):
        invoices = load_invoices()
        for invoice in invoices:
            if invoice['id'] == invoice_id:
                invoice['status'] = new_status
                break
        save_invoices(invoices)
        self.refresh_invoice_list()
        
    def refresh_invoice_list(self):
        self.invoice_table.clear()  # Clear the existing table content
        invoices = load_invoices()
        # Populate the table with invoice data

    def open_settings(self):
        settings_dialog = SettingsDialog(self)
        settings_dialog.exec_()

    def add_item(self):
        item_name = self.item_name_edit.text()
        quantity_text = self.quantity_edit.text()
        price_text = self.price_edit.text()

        try:
            quantity = int(quantity_text)
            price = float(price_text)
            if item_name and quantity > 0 and price > 0:
                total_price = quantity * price
                self.invoice_items.append((item_name, quantity, price, total_price))

                # Add item to the table
                row_position = self.items_table.rowCount()
                self.items_table.insertRow(row_position)
                self.items_table.setItem(row_position, 0, QTableWidgetItem(item_name))
                self.items_table.setItem(row_position, 1, QTableWidgetItem(str(quantity)))
                self.items_table.setItem(row_position, 2, QTableWidgetItem(str(price)))
                self.items_table.setItem(row_position, 3, QTableWidgetItem(str(total_price)))

                # Clear input fields
                self.item_name_edit.clear()
                self.quantity_edit.clear()
                self.price_edit.clear()
            else:
                QMessageBox.warning(self, 'Error', 'Invalid item details. Please enter valid item name, quantity, and price.')
        except ValueError:
            QMessageBox.warning(self, 'Error', 'Quantity and Price must be numbers.')
    
    def collect_invoice_data(self):
        client_name = self.client_name_input.text()  # Assuming QLineEdit for client name
        invoice_amount = float(self.amount_input.text())  # Assuming QLineEdit for amount
        invoice_date = self.date_input.text()  # Assuming QLineEdit or QDateEdit for date
        invoice_data = {
            "client_name": client_name,
            "amount": invoice_amount,
            "date": invoice_date,
            # Include other fields as needed
        }
        return invoice_data
    
    
    def add_invoice(self):
        # Logic to add a new invoice
        new_invoice_id = self.generate_new_invoice_id()
        self.invoices[new_invoice_id] = {"client_name": "", "items": [], "tax_rate": 0, "discount": 0}
        self.invoice_list_widget.addItem(f"Invoice {new_invoice_id}")
        # Open invoice editor or use existing UI to input details

    def edit_invoice(self):
        # Logic to edit selected invoice
        selected_invoice_id = self.get_selected_invoice_id()
        if selected_invoice_id is not None:
            invoice_data = self.invoices[selected_invoice_id]
            # Open invoice editor with this data

    def export_invoice(self):
        # Logic to export selected invoice
        selected_invoice_id = self.get_selected_invoice_id()
        if selected_invoice_id is not None:
            invoice_data = self.invoices[selected_invoice_id]
            # Use existing export logic

    def generate_new_invoice_id(self):
        self.current_invoice_id += 1
        return self.current_invoice_id

    def get_selected_invoice_id(self):
        selected_items = self.invoice_list_widget.selectedItems()
        if selected_items:
            selected_item_text = selected_items[0].text()
            return int(selected_item_text.split()[1])  # Assumes the format "Invoice {id}"
        return None        
    
    def generate_invoice_number(self):
        return f"INV-{datetime.now().strftime('%Y%m%d%H%M%S')}"

    def calculate_total(self):
        # Calculate subtotal
        try:
            subtotal = sum(quantity * price for _, quantity, price, _ in self.invoice_items)
        except ValueError as e:
            QMessageBox.warning(self, 'Error', f'Error calculating subtotal: {e}')
            return None, None

        # Calculate tax and discount
        try:
            tax_rate = float(self.tax_rate_edit.text())
            discount = float(self.discount_edit.text())
        except ValueError:
            QMessageBox.warning(self, 'Error', 'Invalid tax or discount value.')
            return None, None

        tax_amount = subtotal * (tax_rate / 100)
        total = subtotal + tax_amount - discount
        return subtotal, total

    def generate_invoice(self):
        subtotal, total = self.calculate_total()
        if subtotal is not None:
            invoice_summary = f"Subtotal: ${subtotal}\nTax: ${subtotal * (float(self.tax_rate_edit.text()) / 100)}\nDiscount: ${self.discount_edit.text()}\nTotal: ${total}"
            QMessageBox.information(self, 'Invoice Summary', invoice_summary)

    def export_invoice(self):
        client_name = self.client_name_edit.text()
        client_address = self.client_address_edit.text()
        invoice_number = self.generate_invoice_number()
        date = datetime.now().strftime('%Y-%m-%d')
        due_date = (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d')
        subtotal, total = self.calculate_total()

        if subtotal is None:
            return

        # Choose format to export
        format_choice, _ = QInputDialog.getItem(self, "Export Format", "Choose a format:", ["Excel", "Word", "Cancel"], 0, False)
        if format_choice == "Cancel":
            return

        tax_rate = float(self.tax_rate_edit.text())
        discount = float(self.discount_edit.text())
        tax_amount = subtotal * (tax_rate / 100)
        final_total = subtotal + tax_amount - discount

        if format_choice == "Excel":
            export_to_excel(client_name, client_address, self.invoice_items, final_total, invoice_number, date, due_date, tax_amount, discount)
        elif format_choice == "Word":
            export_to_word(client_name, client_address, self.invoice_items, final_total, invoice_number, date, due_date, tax_amount, discount)



class SettingsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Settings')
        layout = QVBoxLayout()

        # Logo selection
        self.logo_path_edit = QLineEdit()
        self.browse_button = QPushButton('Browse')
        self.browse_button.clicked.connect(self.browse_logo)
        layout.addWidget(QLabel('Company Logo:'))
        layout.addWidget(self.logo_path_edit)
        layout.addWidget(self.browse_button)

        # Font selection
        self.font_button = QPushButton('Choose Font')
        self.font_button.clicked.connect(self.choose_font)
        layout.addWidget(self.font_button)

        # Default tax rate and discount
        self.default_tax_rate_edit = QLineEdit()
        self.default_discount_edit = QLineEdit()
        layout.addWidget(QLabel('Default Tax Rate (%):'))
        layout.addWidget(self.default_tax_rate_edit)
        layout.addWidget(QLabel('Default Discount:'))
        layout.addWidget(self.default_discount_edit)

        self.setLayout(layout)

    def browse_logo(self):
        file_name, _ = QFileDialog.getOpenFileName(self, 'Select Logo', '', 'Image files (*.jpg *.png)')
        if file_name:
            self.logo_path_edit.setText(file_name)

    def choose_font(self):
        font, ok = QFontDialog.getFont()
        if ok:
            # Save the selected font for later use
            pass



def export_to_excel(client_name, client_address, items, total, invoice_number, date, due_date, tax_amount, discount):
    import pandas as pd

    # Create DataFrame from items
    df_items = pd.DataFrame(items, columns=['Item Name', 'Quantity', 'Price per Item', 'Total'])
    df_items['Total'] = df_items['Quantity'] * df_items['Price per Item']

    # Add summary rows
    summary_data = [
        {'Item Name': 'Subtotal', 'Total': df_items['Total'].sum()},
        {'Item Name': 'Tax', 'Total': tax_amount},
        {'Item Name': 'Discount', 'Total': -discount},
        {'Item Name': 'Final Total', 'Total': total}
    ]
    df_summary = pd.DataFrame(summary_data)
    final_df = pd.concat([df_items, df_summary], ignore_index=True, sort=False).fillna('')

    # Exporting to Excel
    file_name = f"{invoice_number}_{client_name}_invoice.xlsx"
    final_df.to_excel(file_name, index=False)
    QMessageBox.information(None, 'Export', f'Invoice exported to Excel file {file_name}')

def export_to_word(client_name, client_address, items, total, invoice_number, date, due_date, tax_amount, discount):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading('Invoice', level=1)

    # Client details
    doc.add_paragraph(f"Invoice Number: {invoice_number}\nDate: {date}\nDue Date: {due_date}\n")
    doc.add_paragraph(f"Invoice To:\n{client_name}\n{client_address}\n")

    # Add table for items
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Name'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Price per Item'
    hdr_cells[3].text = 'Total'

    for item in items:
        # Assuming each 'item' is a tuple or list with four elements
        item_name, quantity, price, _ = item  # Adjust unpacking based on your actual data structure
        row_cells = table.add_row().cells
        row_cells[0].text = str(item_name)
        row_cells[1].text = str(quantity)
        row_cells[2].text = str(price)
        row_cells[3].text = str(quantity * price)

    # Summary
    doc.add_paragraph(f"Subtotal: ${sum(quantity * price for _, quantity, price, _ in items)}")
    doc.add_paragraph(f"Tax: ${tax_amount}")
    doc.add_paragraph(f"Discount: ${discount}")
    doc.add_paragraph(f"Total: ${total}")

    # Exporting to Word
    file_name = f"{invoice_number}_{client_name}_invoice.docx"
    doc.save(file_name)
    QMessageBox.information(None, 'Export', f'Invoice exported to Word file {file_name}')

def main():
    app = QApplication(sys.argv)
    ex = InvoiceApp()
    ex.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
